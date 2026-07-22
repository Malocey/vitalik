"""
Multi-Format Extraction Engine für VG Delikatessen.
Unterstützt PDF, PNG, JPG, TXT, MD, DOCX, XLSX, CSV, EML und MSG.
Extrahierte Inhalte werden einheitlich in Seiten-/Abschnitts-Strukturen überführt.
"""

import email
import json
import logging
import os
import re
import hashlib
import zipfile
from email import policy
from pathlib import Path
from typing import Any, Dict, List, Optional

logger = logging.getLogger("MultiFormatEngine")


class MultiFormatEngine:
    MAX_FILE_BYTES = 50 * 1024 * 1024
    MAX_UNCOMPRESSED_BYTES = 200 * 1024 * 1024
    MAX_TABLE_ROWS = 100_000
    SUPPORTED_EXTENSIONS = {
        ".pdf", ".png", ".jpg", ".jpeg", ".txt", ".md",
        ".docx", ".xlsx", ".csv", ".eml"
    }

    def is_supported(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def extract_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """
        Liest eine beliebige unterstützte Datei aus und gibt eine Liste von Seiten-Dicts zurück.
        """
        ext = file_path.suffix.lower()
        file_path = file_path.resolve(strict=True)
        if file_path.stat().st_size > self.MAX_FILE_BYTES:
            raise ValueError(f"Datei überschreitet {self.MAX_FILE_BYTES // (1024 * 1024)} MB")
        if ext == ".pdf":
            from src.parser.pdf_engine import pdf_engine
            return pdf_engine.inspect_pdf(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            from src.parser.ocr_engine import ocr_engine
            ocr_res = ocr_engine.extract_with_quality(file_path)
            text = ocr_res.get("text", "")
            conf = float(ocr_res.get("confidence", 0.0) or 0.0)
            return [{
                "page_num": 1,
                "text_snippet": text[:500],
                "full_text": text,
                "ocr_status": "IMAGE_OCR",
                "ocr_confidence": conf
            }]
        elif ext in [".txt", ".md"]:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return [{
                "page_num": 1,
                "text_snippet": text[:500],
                "full_text": text,
                "ocr_status": "PLAIN_TEXT",
                "ocr_confidence": 1.0
            }]
        elif ext == ".docx":
            return self._extract_docx(file_path)
        elif ext in [".xlsx", ".csv"]:
            return self._extract_spreadsheet(file_path)
        elif ext == ".eml":
            return self._extract_email(file_path)
        else:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
            return [{
                "page_num": 1,
                "text_snippet": text[:500],
                "full_text": text,
                "ocr_status": "UNKNOWN_TEXT",
                "ocr_confidence": 1.0
            }]

    def _extract_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        try:
            self._validate_zip_container(file_path)
            import docx
            doc = docx.Document(str(file_path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Format tables inside docx into Markdown
            for table in doc.tables:
                table_lines = []
                for row in table.rows:
                    cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append("| " + " | ".join(cells) + " |")
                if table_lines:
                    full_text.append("\n".join(table_lines))

            combined = "\n\n".join(full_text)
            return [{
                "page_num": 1,
                "text_snippet": combined[:500],
                "full_text": combined,
                "ocr_status": "DOCX",
                "ocr_confidence": 1.0
            }]
        except Exception as e:
            logger.warning(f"[MultiFormatEngine] Fehler beim Lesen von DOCX {file_path.name}: {e}")
            return [{
                "page_num": 1,
                "text_snippet": f"Fehler beim DOCX-Lesen: {e}",
                "full_text": "",
                "ocr_status": "DOCX_ERROR",
                "ocr_confidence": 0.0
            }]

    def _extract_spreadsheet(self, file_path: Path) -> List[Dict[str, Any]]:
        ext = file_path.suffix.lower()
        pages = []
        try:
            if ext == ".csv":
                import csv
                lines = []
                with open(file_path, "r", encoding="utf-8-sig", errors="replace", newline="") as f:
                    reader = csv.reader(f)
                    for index, row in enumerate(reader):
                        if index >= self.MAX_TABLE_ROWS:
                            raise ValueError("CSV überschreitet das Zeilenlimit")
                        lines.append("| " + " | ".join(row) + " |")
                text = "\n".join(lines)
                pages.append({
                    "page_num": 1,
                    "text_snippet": text[:500],
                    "full_text": text,
                    "ocr_status": "CSV",
                    "ocr_confidence": 1.0
                })
            else:
                import openpyxl
                self._validate_zip_container(file_path)
                wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
                for idx, sheet_name in enumerate(wb.sheetnames, 1):
                    sheet = wb[sheet_name]
                    sheet_lines = [f"=== SHEET: {sheet_name} ==="]
                    for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
                        if row_index >= self.MAX_TABLE_ROWS:
                            raise ValueError("Tabellenblatt überschreitet das Zeilenlimit")
                        if any(row):
                            cells = [str(val) if val is not None else "" for val in row]
                            sheet_lines.append("| " + " | ".join(cells) + " |")
                    text = "\n".join(sheet_lines)
                    pages.append({
                        "page_num": idx,
                        "text_snippet": text[:500],
                        "full_text": text,
                        "ocr_status": "EXCEL",
                        "ocr_confidence": 1.0
                    })
        except Exception as e:
            logger.warning(f"[MultiFormatEngine] Fehler beim Lesen von Excel/CSV {file_path.name}: {e}")
            pages.append({
                "page_num": 1,
                "text_snippet": f"Fehler beim Tabelle-Lesen: {e}",
                "full_text": "",
                "ocr_status": "EXCEL_ERROR",
                "ocr_confidence": 0.0
            })
        return pages

    def _extract_email(self, file_path: Path) -> List[Dict[str, Any]]:
        # Standard .eml parsing
        try:
            raw_bytes = file_path.read_bytes()
            msg = email.message_from_bytes(raw_bytes, policy=policy.default)
            headers = [
                f"From: {msg.get('from', '')}",
                f"To: {msg.get('to', '')}",
                f"Subject: {msg.get('subject', '')}",
                f"Date: {msg.get('date', '')}"
            ]
            body_parts = []
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain":
                        payload = part.get_content()
                        if isinstance(payload, str):
                            body_parts.append(payload)
            else:
                content = msg.get_content() or ""
                body_parts.append(content if msg.get_content_type() == "text/plain" else "")

            combined_body = "\n".join(body_parts)
            text = "\n".join(headers) + "\n\n" + combined_body
            return [{
                "page_num": 1,
                "text_snippet": text[:500],
                "full_text": text,
                "ocr_status": "EMAIL_EML",
                "ocr_confidence": 1.0
            }]
        except Exception as e:
            logger.warning(f"[MultiFormatEngine] Fehler beim Lesen von E-Mail {file_path.name}: {e}")
            return [{
                "page_num": 1,
                "text_snippet": f"Fehler beim E-Mail-Lesen: {e}",
                "full_text": "",
                "ocr_status": "EMAIL_ERROR",
                "ocr_confidence": 0.0
            }]

    def extract_batch_parallel(self, file_paths: List[Path], max_workers: int = 8) -> Dict[str, List[Dict[str, Any]]]:
        """
        Liest mehrere Dateien parallel über einen ThreadPoolExecutor ein.
        """
        from concurrent.futures import ThreadPoolExecutor, as_completed
        results = {}
        if not file_paths:
            return results

        workers = min(max_workers, len(file_paths))
        with ThreadPoolExecutor(max_workers=workers) as executor:
            future_to_file = {executor.submit(self.extract_document, f): str(f) for f in file_paths}
            for future in as_completed(future_to_file):
                f_path = future_to_file[future]
                try:
                    results[f_path] = future.result()
                except Exception as e:
                    logger.warning(f"[MultiFormatEngine] Fehler beim parallelen Lesen von {f_path}: {e}")
                    results[f_path] = [{
                        "page_num": 1,
                        "text_snippet": f"Fehler: {e}",
                        "full_text": "",
                        "ocr_status": "ERROR",
                        "ocr_confidence": 0.0
                    }]
        return results

    def _validate_zip_container(self, file_path: Path) -> None:
        with zipfile.ZipFile(file_path) as archive:
            total_size = sum(item.file_size for item in archive.infolist())
            if total_size > self.MAX_UNCOMPRESSED_BYTES:
                raise ValueError("Office-Datei überschreitet das Entpacklimit")

    @staticmethod
    def _file_md5(file_path: Path) -> str:
        digest = hashlib.md5()
        with file_path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    def persist_non_pdf_document(self, file_path: Path, document: Dict[str, Any]) -> Dict[str, Any]:
        """Persistiert Nicht-PDF-Inhalte ohne sie fälschlich durch den PDF-Sorter zu schicken."""
        from src.core.contact_memory import contact_memory
        from src.core.rag_engine import rag_engine
        from src.wiki.wiki_engine import wiki_engine

        checksum = self._file_md5(file_path)
        page = int(document.get("start_seite") or 1)
        beleg_id = f"MF-{checksum[:16]}-{page}"
        document["md5_hash"] = checksum
        document["beleg_id"] = beleg_id
        document["beleg_link"] = str(file_path.resolve())
        wiki_engine.create_or_update_beleg_page(document, beleg_id)
        rag_engine.index_beleg(document, beleg_id)
        verification = rag_engine.verify_beleg_persistence(beleg_id)
        document["persistence_verification"] = verification
        document["persistence_verified"] = verification["ok"]
        if not verification["ok"]:
            raise RuntimeError(f"Mehrformat-Beleg {beleg_id} nicht vollständig persistent")
        document["contact_memory"] = contact_memory.learn_from_document(document, beleg_id)
        return {"saved_path": str(file_path.resolve()), "beleg_id": beleg_id, "passed": True}


multi_format_engine = MultiFormatEngine()
